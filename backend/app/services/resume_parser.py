"""
TalentMatch AI - Resume Parser Service
Extracts structured information from PDF, DOC, and DOCX resumes
Uses PyMuPDF, python-docx, spaCy NLP, and Regex
"""

import re
import logging
from pathlib import Path
from typing import Optional, List, Dict, Any, Tuple
import json

logger = logging.getLogger(__name__)

# ─────────────────────────────────────────
# Known skills database for matching
# ─────────────────────────────────────────
TECHNICAL_SKILLS = {
    # Languages
    "python", "javascript", "typescript", "java", "c++", "c#", "go", "rust",
    "ruby", "php", "swift", "kotlin", "scala", "r", "matlab", "perl",
    "bash", "shell", "powershell", "sql", "html", "css", "xml", "yaml",

    # Frontend
    "react", "vue", "angular", "next.js", "nuxt", "svelte", "redux",
    "jquery", "bootstrap", "tailwind", "sass", "less", "webpack", "vite",
    "graphql", "apollo", "react native", "flutter", "ionic",

    # Backend
    "node.js", "express", "fastapi", "django", "flask", "spring", "rails",
    "laravel", "nestjs", "fastify", "gin", "fiber", "actix", "asp.net",

    # Databases
    "postgresql", "mysql", "mongodb", "redis", "elasticsearch",
    "cassandra", "dynamodb", "sqlite", "oracle", "mariadb", "neo4j",
    "firebase", "supabase",

    # Cloud & DevOps
    "aws", "gcp", "azure", "docker", "kubernetes", "terraform", "ansible",
    "jenkins", "github actions", "gitlab ci", "circleci", "nginx", "apache",
    "linux", "unix",

    # AI/ML
    "machine learning", "deep learning", "tensorflow", "pytorch", "keras",
    "scikit-learn", "pandas", "numpy", "opencv", "nlp", "computer vision",
    "llm", "openai", "langchain", "huggingface", "transformers",

    # Tools
    "git", "github", "gitlab", "bitbucket", "jira", "confluence",
    "figma", "postman", "swagger", "kafka", "rabbitmq", "celery",
}

SOFT_SKILLS = {
    "leadership", "communication", "teamwork", "problem solving",
    "critical thinking", "project management", "agile", "scrum",
    "collaboration", "adaptability", "creativity", "time management",
    "mentoring", "analytical", "attention to detail", "decision making",
}

EDUCATION_LEVELS = {
    "phd": 5, "ph.d": 5, "doctorate": 5,
    "master": 4, "msc": 4, "mba": 4, "m.s.": 4, "m.tech": 4,
    "bachelor": 3, "bsc": 3, "b.s.": 3, "b.tech": 3, "b.e.": 3, "undergraduate": 3,
    "associate": 2, "diploma": 2,
    "high school": 1, "secondary": 1,
}

# Regex patterns
EMAIL_PATTERN = re.compile(
    r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
)
PHONE_PATTERN = re.compile(
    r'[\+]?[(]?[0-9]{3}[)]?[-\s\.]?[0-9]{3}[-\s\.]?[0-9]{4,6}'
)
LINKEDIN_PATTERN = re.compile(
    r'linkedin\.com/in/[\w\-]+'
)
GITHUB_PATTERN = re.compile(
    r'github\.com/[\w\-]+'
)
YEARS_EXP_PATTERN = re.compile(
    r'(\d+(?:\.\d+)?)\s*\+?\s*(?:years?|yrs?)\s*(?:of\s+)?(?:experience|exp)',
    re.IGNORECASE
)


class ResumeParser:
    """
    Multi-format resume parser that extracts structured candidate information.
    """

    def extract_text(self, file_path: str) -> str:
        """
        Extract raw text from PDF, DOC, or DOCX file.
        """
        path = Path(file_path)
        ext = path.suffix.lower()

        if ext == ".pdf":
            return self._extract_pdf(file_path)
        elif ext in (".docx",):
            return self._extract_docx(file_path)
        elif ext == ".doc":
            return self._extract_doc(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")

    def _extract_pdf(self, file_path: str) -> str:
        """Extract text from PDF using PyMuPDF (fitz)."""
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(file_path)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text.strip()
        except ImportError:
            logger.error("PyMuPDF not installed. Run: pip install pymupdf")
            raise
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            return ""

    def _extract_docx(self, file_path: str) -> str:
        """Extract text from DOCX using python-docx."""
        try:
            from docx import Document
            doc = Document(file_path)
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text)

            return "\n".join(paragraphs).strip()
        except ImportError:
            logger.error("python-docx not installed. Run: pip install python-docx")
            raise
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            return ""

    def _extract_doc(self, file_path: str) -> str:
        """Extract text from legacy .doc file (limited support)."""
        try:
            # Try using antiword or python-docx2txt as fallback
            import docx2txt
            return docx2txt.process(file_path)
        except ImportError:
            logger.warning("docx2txt not installed. .doc support limited.")
            return ""
        except Exception as e:
            logger.error(f"DOC extraction failed: {e}")
            return ""

    def parse(self, file_path: str, resume_text: Optional[str] = None) -> Dict[str, Any]:
        """
        Main parse method - extracts all candidate information.

        Returns a structured dict with all parsed fields.
        """
        if resume_text is None:
            resume_text = self.extract_text(file_path)

        result = {
            "resume_text": resume_text,
            "name": self._extract_name(resume_text),
            "email": self._extract_email(resume_text),
            "phone": self._extract_phone(resume_text),
            "linkedin_url": self._extract_linkedin(resume_text),
            "github_url": self._extract_github(resume_text),
            "location": self._extract_location(resume_text),
            "skills": self._extract_skills(resume_text),
            "soft_skills": self._extract_soft_skills(resume_text),
            "education_details": self._extract_education(resume_text),
            "education_level": None,
            "work_experience": self._extract_experience(resume_text),
            "years_of_experience": self._extract_years_of_experience(resume_text),
            "projects": self._extract_projects(resume_text),
            "certifications": self._extract_certifications(resume_text),
            "technologies": [],
            "languages": [],
        }

        # Derive education level from education details
        result["education_level"] = self._derive_education_level(result["education_details"])

        # Separate languages from general skills
        lang_skills = [s for s in result["skills"] if s in {
            "python", "javascript", "typescript", "java", "c++", "c#",
            "go", "rust", "ruby", "php", "swift", "kotlin", "scala",
            "r", "matlab", "perl", "bash", "sql", "html", "css"
        }]
        result["languages"] = lang_skills

        # Collect all technologies
        result["technologies"] = list(set(result["skills"] + [
            t for exp in result["work_experience"]
            for t in exp.get("technologies", [])
        ] + [
            t for proj in result["projects"]
            for t in proj.get("technologies", [])
        ]))

        return result

    # ─────────────────────────────────────────
    # Field Extractors
    # ─────────────────────────────────────────

    def _extract_name(self, text: str) -> str:
        """Extract candidate name from resume header."""
        lines = text.strip().split('\n')

        # First non-empty line is usually the name
        for line in lines[:5]:
            line = line.strip()
            if not line:
                continue
            # Filter out common header words
            skip_words = {'resume', 'curriculum', 'vitae', 'cv', 'profile', 'page'}
            if any(w in line.lower() for w in skip_words):
                continue
            # Name lines are typically short, no digits, no special chars
            if len(line) < 60 and not re.search(r'[\d@#$%]', line):
                # Try to validate it looks like a name (2+ words or 1 word capitalized)
                words = line.split()
                if 1 <= len(words) <= 5 and all(w[0].isupper() for w in words if w):
                    return line

        # Fallback: try spaCy NER
        try:
            import spacy
            nlp = spacy.load("en_core_web_sm")
            doc = nlp(text[:500])
            for ent in doc.ents:
                if ent.label_ == "PERSON":
                    return ent.text
        except Exception:
            pass

        return "Unknown Candidate"

    def _extract_email(self, text: str) -> Optional[str]:
        """Extract email address."""
        match = EMAIL_PATTERN.search(text)
        return match.group(0) if match else None

    def _extract_phone(self, text: str) -> Optional[str]:
        """Extract phone number."""
        match = PHONE_PATTERN.search(text)
        return match.group(0).strip() if match else None

    def _extract_linkedin(self, text: str) -> Optional[str]:
        """Extract LinkedIn URL."""
        match = LINKEDIN_PATTERN.search(text)
        if match:
            return f"https://www.{match.group(0)}"
        return None

    def _extract_github(self, text: str) -> Optional[str]:
        """Extract GitHub URL."""
        match = GITHUB_PATTERN.search(text)
        if match:
            return f"https://www.{match.group(0)}"
        return None

    def _extract_location(self, text: str) -> Optional[str]:
        """Extract location/city information."""
        # Common patterns: "City, State" or "City, Country"
        location_pattern = re.compile(
            r'\b([A-Z][a-zA-Z\s]+),\s*([A-Z]{2}|[A-Z][a-zA-Z\s]+)\b'
        )
        # Look in first 300 chars (usually in header)
        match = location_pattern.search(text[:300])
        if match:
            return match.group(0)

        # Try spaCy GPE entities
        try:
            import spacy
            nlp = spacy.load("en_core_web_sm")
            doc = nlp(text[:500])
            for ent in doc.ents:
                if ent.label_ in ("GPE", "LOC"):
                    return ent.text
        except Exception:
            pass

        return None

    def _extract_skills(self, text: str) -> List[str]:
        """Extract technical skills from resume text."""
        text_lower = text.lower()
        found_skills = []

        for skill in TECHNICAL_SKILLS:
            # Use word boundary matching for short skills to avoid false positives
            pattern = r'\b' + re.escape(skill) + r'\b'
            if re.search(pattern, text_lower):
                found_skills.append(skill)

        # Also look for skills section
        skills_section = self._extract_section(text, ["skills", "technical skills", "core competencies", "expertise"])
        if skills_section:
            # Parse comma/bullet-separated skills
            raw_skills = re.split(r'[,•|\n\t]', skills_section)
            for s in raw_skills:
                s_clean = s.strip().lower()
                if 2 < len(s_clean) < 50:
                    found_skills.append(s_clean)

        return list(set(found_skills))

    def _extract_soft_skills(self, text: str) -> List[str]:
        """Extract soft skills from resume."""
        text_lower = text.lower()
        found = []
        for skill in SOFT_SKILLS:
            if skill in text_lower:
                found.append(skill)
        return list(set(found))

    def _extract_education(self, text: str) -> List[Dict[str, Any]]:
        """Extract education details."""
        education = []
        section = self._extract_section(text, ["education", "academic background", "qualifications"])

        if not section:
            section = text

        # Common degree patterns
        degree_patterns = [
            r'(Ph\.?D\.?|Doctor(?:ate)?)\s+(?:of|in|–)?\s*([\w\s]+)',
            r'(M\.?S\.?|M\.?Sc\.?|Master(?:s|\'s)?(?:\s+of\s+Science)?)\s+(?:of|in|–)?\s*([\w\s]+)',
            r'(MBA|Master(?:s|\'s)?(?:\s+of\s+Business\s+Administration))',
            r'(B\.?S\.?|B\.?Sc\.?|B\.?E\.?|B\.?Tech\.?|Bachelor(?:s|\'s)?)\s+(?:of|in|–)?\s*([\w\s]+)',
        ]

        for pattern in degree_patterns:
            matches = re.finditer(pattern, section, re.IGNORECASE)
            for match in matches:
                degree = match.group(1) if match.lastindex >= 1 else ""
                field = match.group(2).strip() if match.lastindex >= 2 else ""

                # Try to find institution near this match
                context = section[max(0, match.start()-200):match.end()+200]
                year_match = re.search(r'\b(19|20)\d{2}\b', context)
                year = year_match.group(0) if year_match else None

                education.append({
                    "degree": degree,
                    "field": field[:100],
                    "institution": None,  # TODO: NER-based institution extraction
                    "year": year,
                    "gpa": None,
                })

        return education[:5]  # Return max 5 education entries

    def _derive_education_level(self, education_details: List[Dict]) -> Optional[str]:
        """Determine highest education level from parsed education."""
        highest_level = 0
        highest_label = None

        for edu in education_details:
            degree = (edu.get("degree") or "").lower()
            for label, level in EDUCATION_LEVELS.items():
                if label in degree and level > highest_level:
                    highest_level = level
                    highest_label = edu.get("degree")

        return highest_label

    def _extract_experience(self, text: str) -> List[Dict[str, Any]]:
        """Extract work experience entries."""
        experience = []
        section = self._extract_section(text, [
            "experience", "work experience", "professional experience",
            "employment", "work history", "career history"
        ])

        if not section:
            return []

        # Split by job patterns (typically company name + dates)
        # This is a simplified heuristic approach
        blocks = re.split(r'\n(?=[A-Z][^\n]{5,50}\n)', section)

        for block in blocks[:10]:  # Max 10 experience entries
            if not block.strip():
                continue

            lines = [l.strip() for l in block.split('\n') if l.strip()]
            if not lines:
                continue

            # Extract date range
            date_pattern = re.compile(
                r'((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[\w]*\.?\s+\d{4}|'
                r'\d{4})\s*[-–—to]+\s*((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[\w]*\.?\s+\d{4}|'
                r'\d{4}|Present|Current|Now)',
                re.IGNORECASE
            )
            date_match = date_pattern.search(block)
            start_date = date_match.group(1) if date_match else None
            end_date = date_match.group(2) if date_match else None

            # Extract technologies mentioned in block
            block_lower = block.lower()
            techs = [s for s in TECHNICAL_SKILLS if re.search(r'\b' + re.escape(s) + r'\b', block_lower)]

            exp_entry = {
                "title": lines[0] if lines else None,
                "company": lines[1] if len(lines) > 1 else None,
                "start_date": start_date,
                "end_date": end_date,
                "duration": f"{start_date} - {end_date}" if start_date else None,
                "description": " ".join(lines[2:5]) if len(lines) > 2 else None,
                "technologies": techs[:10],
            }
            experience.append(exp_entry)

        return experience

    def _extract_years_of_experience(self, text: str) -> float:
        """Extract or estimate years of work experience."""
        # Direct mention pattern
        match = YEARS_EXP_PATTERN.search(text)
        if match:
            return float(match.group(1))

        # Count unique years in experience section
        years = re.findall(r'\b(20\d{2}|19\d{2})\b', text)
        if years:
            years_int = [int(y) for y in years]
            span = max(years_int) - min(years_int)
            return float(min(span, 25))  # Cap at 25 years

        return 0.0

    def _extract_projects(self, text: str) -> List[Dict[str, Any]]:
        """Extract project information."""
        projects = []
        section = self._extract_section(text, ["projects", "personal projects", "key projects", "portfolio"])

        if not section:
            return []

        # Split by project name patterns
        blocks = re.split(r'\n(?=[A-Z•\-])', section)

        for block in blocks[:10]:
            if not block.strip() or len(block) < 20:
                continue

            lines = [l.strip() for l in block.split('\n') if l.strip()]
            block_lower = block.lower()

            # Extract technologies
            techs = [s for s in TECHNICAL_SKILLS if re.search(r'\b' + re.escape(s) + r'\b', block_lower)]

            # Extract URL
            url_match = re.search(r'https?://\S+', block)

            projects.append({
                "name": lines[0][:100] if lines else None,
                "description": " ".join(lines[1:3])[:300] if len(lines) > 1 else None,
                "technologies": techs[:8],
                "url": url_match.group(0) if url_match else None,
            })

        return projects

    def _extract_certifications(self, text: str) -> List[str]:
        """Extract certifications."""
        section = self._extract_section(text, [
            "certifications", "certificates", "licenses", "credentials", "awards"
        ])

        if not section:
            return []

        # Split by newlines and bullets
        certs = re.split(r'[•\n\-]', section)
        return [c.strip() for c in certs if len(c.strip()) > 5][:10]

    def _extract_section(self, text: str, section_names: List[str]) -> Optional[str]:
        """
        Extract a specific section from resume text.
        Looks for section headers and returns content until the next section.
        """
        lines = text.split('\n')
        section_start = None
        section_content = []

        common_headers = {
            "experience", "education", "skills", "projects", "certifications",
            "summary", "objective", "awards", "publications", "references",
            "volunteer", "activities", "languages", "interests", "contact"
        }

        for i, line in enumerate(lines):
            line_lower = line.strip().lower()

            # Check if this line is our target section header
            if section_start is None:
                for name in section_names:
                    if line_lower == name or line_lower.startswith(name + ':'):
                        section_start = i
                        break
            else:
                # Check if we've hit another section header
                is_new_section = any(
                    line_lower == h or line_lower.startswith(h + ':')
                    for h in common_headers
                    if h not in section_names
                )
                if is_new_section and len(section_content) > 2:
                    break
                section_content.append(line)

        if section_content:
            return "\n".join(section_content[:100])  # Limit to 100 lines

        return None


# Singleton instance
resume_parser = ResumeParser()
